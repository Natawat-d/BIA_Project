"""Convert docs/FINAL_REPORT.md -> docs/Final_Report.docx (AIT thesis-template styling).

Layout/typography mirrors ait-thesis-latex-template (aitthesis.cls + aitfont12.sty):
  - A4 paper (the class forces a4paper; only 12pt is supported).
  - Margins: left 1.25" (1in + \\oddsidemargin .25in), right 0.92"
    (A4 8.27" - 1.25" - \\textwidth 6.1"), top 0.95" (1in + \\topmargin -.8in
    + \\headheight 12pt + \\headsep 42pt), bottom 1.25" (11.69" - 0.95"
    - \\textheight 9.5in).
  - Times New Roman 12; body = \\singlespace (exact 14pt baselineskip),
    \\parskip 12pt between paragraphs, \\parindent 0, justified.
  - H1 = AIT chapter head: new page, centered bold 14pt (\\large); numbered
    chapters render as "Chapter N" + title on the next line (not uppercase).
  - H2/H3 = \\section/\\subsection: 12pt bold flush left (3.5ex/2.3ex and
    3.25ex/1.5ex skips ~ 18/12pt and 17/8pt).
  - Table captions ("Table N.N — ..." directly above a table): \\small (11pt),
    bold "Table N.N" label (\\fnum@table), centered; table cells single-spaced.
  - Page number bottom center (\\pagestyle{plain}, used for the thesis body).
  - Title block (before the first ---) centered, title line bold (aittitle.sty).

Run inside the dashboard container:
  pip install python-docx
  python /app/reports/make_final_report_docx.py
Handles the markdown subset used by the report: #/##/### headings, tables,
bullet lists, numbered lists (literal numbers preserved), **bold**, *italic*,
`code`, and --- rules (first rule = end of the title block; each H1 starts a
new page, so no explicit break is emitted).
"""
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

SRC = Path("/app/docs/FINAL_REPORT.md")
OUT = Path("/app/docs/Final_Report.docx")

doc = Document()

# --- page setup: A4 + AIT margins (derived from aitfont12.sty) --------------
for sec in doc.sections:
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.left_margin = Inches(1.25)   # 1in + \oddsidemargin .25in
    sec.right_margin = Inches(0.92)  # 8.27 - 1.25 - \textwidth 6.1in
    sec.top_margin = Inches(0.95)    # 1in - .8in + headheight 12pt + headsep 42pt
    sec.bottom_margin = Inches(1.25)  # 11.69 - 0.95 - \textheight 9.5in
    sec.header_distance = Inches(0.2)
    sec.footer_distance = Inches(0.9)

# --- base styles -------------------------------------------------------------
BLACK = RGBColor(0, 0, 0)
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = Pt(14)  # \singlespace: baselineskip=14pt
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(12)   # \parskip 12pt

for name, size, before, after in (
    ("Heading 1", 14, 0, 20),   # chapter head: \large bold, \vskip 20pt after
    ("Heading 2", 12, 18, 12),  # \section: \normalsize\bf, 3.5ex / 2.3ex
    ("Heading 3", 12, 17, 8),   # \subsection: \normalsize\bf, 3.25ex / 1.5ex
):
    st = doc.styles[name]
    st.font.name = "Times New Roman"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.italic = False
    st.font.color.rgb = BLACK
    pf = st.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.0
    pf.keep_with_next = True
doc.styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

# --- page number bottom center (\pagestyle{plain}) ---------------------------
foot_p = doc.sections[0].footer.paragraphs[0]
foot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
foot_p.paragraph_format.space_after = Pt(0)
foot_p.paragraph_format.line_spacing = 1.0
fld_run = foot_p.add_run()
for tag, attrs, text in (
    ("w:fldChar", {"w:fldCharType": "begin"}, None),
    ("w:instrText", {"xml:space": "preserve"}, "PAGE"),
    ("w:fldChar", {"w:fldCharType": "end"}, None),
):
    el = OxmlElement(tag)
    for k, v in attrs.items():
        el.set(qn(k), v)
    if text:
        el.text = text
    fld_run._r.append(el)

props = doc.core_properties
props.title = "Netflix Customer Retention Intelligence DSS — Final Report"
props.author = "Natawat Damrongsilp; Liza Shrestha; Subhana Chitrakar"


INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")
CHAPNUM = re.compile(r"^(\d+)[.:]?\s+(.+)$")
CAPTION = re.compile(r"^((?:Table|Figure)\s+\d+(?:\.\d+)?)(.*)$")


def add_runs(par, text):
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            par.add_run(tok[2:-2]).bold = True
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            par.add_run(tok[1:-1]).italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = par.add_run(tok[1:-1])
            r.font.name = "Courier New"  # \ttdefault -> cmtt (typewriter)
            r.font.size = Pt(11)
        else:
            par.add_run(tok)


def add_chapter_head(text):
    """AIT \\chapter: new page, centered bold 14pt; numbered chapters render
    as 'Chapter N' over the title (uppercase-ing is commented out in the cls)."""
    m = CHAPNUM.match(text)
    first = doc.add_heading(f"Chapter {m.group(1)}" if m else text, level=1)
    # 'page break before' is a no-op when the paragraph already tops a page
    first.paragraph_format.page_break_before = True
    if m:
        first.paragraph_format.space_after = Pt(6)  # \vskip -6pt + \vspace{1em}
        doc.add_heading(m.group(2), level=1)


def add_caption(text):
    """\\fnum@table: \\small bold 'Table N.N' label, centered above the table."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_after = Pt(6)
    m = CAPTION.match(text)
    label = p.add_run(m.group(1))
    label.bold = True
    add_runs(p, m.group(2))
    for r in p.runs:
        r.font.size = Pt(11)  # \small


lines = SRC.read_text().split("\n")


def next_nonblank(idx):
    j = idx + 1
    while j < len(lines) and lines[j].strip() == "":
        j += 1
    return lines[j].strip() if j < len(lines) else ""


i, seen_rule, title_done = 0, False, False
while i < len(lines):
    s = lines[i].strip()

    # table
    if s.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|?$", lines[i + 1].strip()):
        header = [c.strip() for c in s.strip("|").split("|")]
        rows = []
        j = i + 2
        while j < len(lines) and lines[j].strip().startswith("|"):
            cells = [c.strip() for c in lines[j].strip().strip("|").split("|")]
            rows.append((cells + [""] * len(header))[: len(header)])
            j += 1
        table = doc.add_table(rows=1 + len(rows), cols=len(header))
        table.style = "Table Grid"
        for k, h in enumerate(header):
            p = table.rows[0].cells[k].paragraphs[0]
            add_runs(p, h)
            for r in p.runs:
                r.bold = True
        for ri, row in enumerate(rows):
            for k, cell in enumerate(row):
                add_runs(table.rows[ri + 1].cells[k].paragraphs[0], cell)
        for row in table.rows:  # tables are single-spaced (sspace) in the class
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.0
        doc.add_paragraph()
        i = j
        continue

    if s == "---":
        seen_rule = True  # end of title block; H1s carry their own page breaks
        i += 1
        continue
    if s == "":
        i += 1
        continue
    if s.startswith("### "):
        doc.add_heading(s[4:], level=3)
    elif s.startswith("## "):
        doc.add_heading(s[3:], level=2)
    elif s.startswith("# "):
        add_chapter_head(s[2:])
    elif s.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs(p, s[2:])
    elif re.match(r"^\d+\. ", s):
        p = doc.add_paragraph()  # keep literal numbers -> no docx renumbering bugs
        p.paragraph_format.left_indent = Inches(0.3)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs(p, s)
    elif CAPTION.match(s) and next_nonblank(i).startswith("|"):
        add_caption(s)
    elif not seen_rule:
        # title block: aittitle.sty centers everything, title line bold
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_runs(p, s)
        if not title_done:
            for r in p.runs:
                r.bold = True
            title_done = True
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs(p, s)
    i += 1

doc.save(OUT)
print(f"wrote {OUT}")
